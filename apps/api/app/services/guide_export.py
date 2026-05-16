"""DOCX export service for generated guides.

Converts a validated Guide Pydantic model into a Word .docx document and
returns the raw bytes. No images are embedded in this first version; steps
that have associated frame screenshots get a plain-text note instead.
"""
from __future__ import annotations

import re
import uuid
from io import BytesIO
from typing import Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.schemas.guide import Guide


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------

def safe_filename(title: str, session_id: Union[uuid.UUID, str]) -> str:
    """Return a safe .docx filename derived from *title*.

    Falls back to ``guide-{session_id}.docx`` when *title* is empty or
    produces an empty slug after sanitisation.
    """
    slug = re.sub(r"[^\w\s-]", "", title or "").strip()
    slug = re.sub(r"[\s_]+", "-", slug)[:60].strip("-")
    if not slug:
        return f"guide-{session_id}.docx"
    return f"{slug}.docx"


def guide_to_docx(
    guide: Guide,
    session_id: Union[uuid.UUID, str],
) -> bytes:
    """Convert *guide* to a .docx document and return the raw bytes.

    All fields are treated defensively: missing or empty optional fields are
    simply skipped rather than raising an error.
    """
    doc = Document()

    # -----------------------------------------------------------------------
    # Remove the default empty paragraph that Word adds to new documents
    # -----------------------------------------------------------------------
    _remove_first_empty_paragraph(doc)

    # -----------------------------------------------------------------------
    # Title
    # -----------------------------------------------------------------------
    doc.add_heading(guide.title or "Guide", level=1)

    # -----------------------------------------------------------------------
    # Metadata block (small, muted paragraph)
    # -----------------------------------------------------------------------
    meta = guide.metadata
    meta_parts: list[str] = []
    if str(session_id):
        meta_parts.append(f"Session: {session_id}")
    if meta:
        if meta.generated_at:
            meta_parts.append(f"Generated: {meta.generated_at}")
        if meta.generated_by:
            meta_parts.append(f"Model: {meta.generated_by}")
        if meta.source_duration_sec is not None:
            meta_parts.append(f"Recording: {_fmt_duration(meta.source_duration_sec)}")
    if meta_parts:
        p = doc.add_paragraph(" · ".join(meta_parts))
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.runs[0].italic = True

    # -----------------------------------------------------------------------
    # Overview
    # -----------------------------------------------------------------------
    if guide.summary:
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(guide.summary)

    # -----------------------------------------------------------------------
    # Estimated duration
    # -----------------------------------------------------------------------
    if guide.estimated_duration_minutes is not None:
        doc.add_heading("Estimated Duration", level=2)
        doc.add_paragraph(f"~{guide.estimated_duration_minutes:.0f} minutes")

    # -----------------------------------------------------------------------
    # Prerequisites & tools
    # -----------------------------------------------------------------------
    has_prereqs = bool(guide.prerequisites)
    has_tools = bool(guide.tools_or_systems)
    if has_prereqs or has_tools:
        doc.add_heading("Prerequisites", level=2)
        if has_prereqs:
            for item in guide.prerequisites:
                doc.add_paragraph(item, style="List Bullet")
        if has_tools:
            doc.add_heading("Tools & Systems", level=3)
            for item in guide.tools_or_systems:
                doc.add_paragraph(item, style="List Bullet")

    # -----------------------------------------------------------------------
    # Global warnings
    # -----------------------------------------------------------------------
    if guide.warnings:
        doc.add_heading("Warnings", level=2)
        for w in guide.warnings:
            _add_callout(doc, f"⚠  {w}")

    # -----------------------------------------------------------------------
    # Global notes
    # -----------------------------------------------------------------------
    if guide.notes:
        doc.add_heading("Notes", level=2)
        for n in guide.notes:
            _add_callout(doc, f"ℹ  {n}")

    # -----------------------------------------------------------------------
    # Procedure — numbered steps
    # -----------------------------------------------------------------------
    if guide.steps:
        doc.add_heading("Procedure", level=2)
        for step in sorted(guide.steps, key=lambda s: s.order):
            step_num = step.order
            doc.add_heading(f"Step {step_num} – {step.title}", level=3)

            if step.description:
                doc.add_paragraph(step.description)

            # Actions
            if step.actions:
                p = doc.add_paragraph()
                run = p.add_run("Actions")
                run.bold = True
                run.font.size = Pt(10)
                for action in step.actions:
                    parts = [action.verb, action.target]
                    if action.value:
                        parts.append(f"→ {action.value}")
                    doc.add_paragraph(" ".join(parts), style="List Bullet")

            # Step-level warnings
            for w in (step.warnings or []):
                _add_callout(doc, f"⚠  {w}")

            # Step-level notes
            for n in (step.notes or []):
                _add_callout(doc, f"ℹ  {n}")

            # Screenshot placeholder
            frame_count = len(step.evidence.frame_keys) if step.evidence else 0
            if frame_count:
                p = doc.add_paragraph(
                    f"[{frame_count} screenshot{'s' if frame_count != 1 else ''} "
                    "available in the online guide]"
                )
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # -----------------------------------------------------------------------
    # Troubleshooting table
    # -----------------------------------------------------------------------
    if guide.troubleshooting:
        doc.add_heading("Troubleshooting", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ("Symptom", "Likely Cause", "Resolution")):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        for entry in guide.troubleshooting:
            row = table.add_row().cells
            row[0].text = entry.symptom or ""
            row[1].text = entry.likely_cause or ""
            row[2].text = entry.resolution or ""

    # -----------------------------------------------------------------------
    # Serialise to bytes
    # -----------------------------------------------------------------------
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _remove_first_empty_paragraph(doc: Document) -> None:
    """Remove the initial empty paragraph that python-docx adds by default."""
    try:
        first = doc.paragraphs[0]
        if not first.text:
            p = first._element
            p.getparent().remove(p)
    except (IndexError, AttributeError):
        pass


def _fmt_duration(seconds: float) -> str:
    seconds = int(seconds)
    if seconds < 60:
        return f"{seconds}s"
    m, s = divmod(seconds, 60)
    return f"{m}m {s}s" if s else f"{m}m"


def _add_callout(doc: Document, text: str) -> None:
    """Add a lightly-styled paragraph used for warnings and notes."""
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Pt(18)
    if p.runs:
        p.runs[0].font.size = Pt(10)
