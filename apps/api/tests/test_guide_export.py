"""Tests for the guide_export service and the export.docx endpoint.

Unit tests for `guide_to_docx` and `safe_filename` run without a database.
Endpoint tests use FastAPI's TestClient with a monkeypatched session service
so they also require no real database.
"""
from __future__ import annotations

import uuid
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.schemas.guide import (
    Action,
    Evidence,
    Guide,
    GuideMetadata,
    Step,
    Troubleshooting,
)
from app.services.guide_export import guide_to_docx, safe_filename


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _make_guide() -> Guide:
    return Guide(
        title="How to configure the firewall",
        summary="This guide explains the steps to configure the company firewall.",
        estimated_duration_minutes=10,
        prerequisites=["Admin access", "VPN connected"],
        tools_or_systems=["pfSense dashboard"],
        warnings=["Do not change rules during business hours"],
        notes=["Back up the current rule set first"],
        steps=[
            Step(
                id="s1",
                order=1,
                title="Open the admin panel",
                description="Navigate to the pfSense dashboard login page.",
                actions=[
                    Action(verb="Open", target="browser", value="https://pfsense.local"),
                    Action(verb="Enter", target="username", value="admin"),
                ],
                evidence=Evidence(frame_keys=["sessions/abc/frames/f001.jpg"]),
                warnings=["Ensure you are on the internal network"],
                notes=["Use Firefox for best compatibility"],
                confidence=0.9,
            ),
            Step(
                id="s2",
                order=2,
                title="Navigate to Firewall Rules",
                description="Click Firewall > Rules in the top menu.",
                actions=[Action(verb="Click", target="Firewall > Rules")],
                evidence=Evidence(frame_keys=[]),
                confidence=0.85,
            ),
        ],
        troubleshooting=[
            Troubleshooting(
                symptom="Login page not loading",
                likely_cause="VPN not connected",
                resolution="Connect to VPN and retry",
            )
        ],
        metadata=GuideMetadata(
            generated_by="gpt-4o",
            generated_at="2026-05-15T10:00:00Z",
            source_session_id="test-session",
            source_duration_sec=300.0,
        ),
    )


# ---------------------------------------------------------------------------
# Unit tests — guide_to_docx
# ---------------------------------------------------------------------------

class TestGuideToDOCX:
    def test_returns_bytes(self):
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        assert isinstance(result, bytes)

    def test_output_is_valid_zip(self):
        """A .docx file is a ZIP archive; the first two bytes must be b'PK'."""
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        assert result[:2] == b"PK"

    def test_document_contains_title(self):
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "How to configure the firewall" in full_text

    def test_document_contains_summary(self):
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert guide.summary in full_text

    def test_document_contains_step_title(self):
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Open the admin panel" in full_text

    def test_document_contains_action(self):
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Action: "Open browser → https://pfsense.local"
        assert "Open" in full_text
        assert "browser" in full_text

    def test_document_contains_troubleshooting(self):
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        doc = Document(BytesIO(result))
        # Tables are not in doc.paragraphs; check table cells
        all_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "Login page not loading" in table_text
        assert "VPN not connected" in table_text

    def test_frame_keys_produce_placeholder_text(self):
        """Step with frame_keys should emit a screenshot placeholder, not crash."""
        guide = _make_guide()
        result = guide_to_docx(guide, uuid.uuid4())
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "screenshot" in full_text.lower()
        assert "1 screenshot" in full_text

    def test_empty_optional_fields_do_not_crash(self):
        """Minimal guide with only required fields must not raise."""
        guide = Guide(
            title="Minimal",
            summary="",
            steps=[
                Step(
                    id="s1",
                    order=1,
                    title="Do something",
                    description="",
                    evidence=Evidence(),
                    confidence=0.5,
                )
            ],
        )
        result = guide_to_docx(guide, "test-id")
        assert result[:2] == b"PK"

    def test_no_steps_does_not_crash(self):
        guide = Guide(title="Empty guide", summary="Nothing here", steps=[])
        result = guide_to_docx(guide, uuid.uuid4())
        assert result[:2] == b"PK"


# ---------------------------------------------------------------------------
# Unit tests — safe_filename
# ---------------------------------------------------------------------------

class TestSafeFilename:
    def test_normal_title(self):
        assert safe_filename("Configure the Firewall", "abc") == "Configure-the-Firewall.docx"

    def test_special_characters_removed(self):
        name = safe_filename("Guide: Setup & Deploy!", "abc")
        assert ":" not in name
        assert "&" not in name
        assert "!" not in name
        assert name.endswith(".docx")

    def test_empty_title_uses_fallback(self):
        sid = "my-session-id"
        assert safe_filename("", sid) == f"guide-{sid}.docx"

    def test_whitespace_only_title_uses_fallback(self):
        sid = "my-session-id"
        assert safe_filename("   ", sid) == f"guide-{sid}.docx"

    def test_long_title_is_truncated(self):
        long_title = "A" * 200
        result = safe_filename(long_title, "sid")
        # slug ≤ 60 chars + ".docx"
        assert len(result) <= 65


# ---------------------------------------------------------------------------
# Endpoint integration tests (no real DB)
# ---------------------------------------------------------------------------

SESSION_ID = uuid.UUID("00000000-0000-0000-0000-000000000001")


def _make_session_with_guide():
    guide = _make_guide()
    return SimpleNamespace(
        id=SESSION_ID,
        title="Firewall guide",
        status="ready",
        guide_content=guide.model_dump(),
    )


def _make_session_without_guide():
    return SimpleNamespace(
        id=SESSION_ID,
        title="Incomplete",
        status="processing",
        guide_content=None,
    )


class TestExportEndpoint:
    def test_returns_200_and_docx_content_type(self):
        session_obj = _make_session_with_guide()
        with patch(
            "app.services.session_service.get_session",
            new=AsyncMock(return_value=session_obj),
        ):
            client = TestClient(app, raise_server_exceptions=False)
            resp = client.get(f"/api/v1/sessions/{SESSION_ID}/export.docx")

        assert resp.status_code == 200
        assert (
            "wordprocessingml.document" in resp.headers.get("content-type", "")
        )

    def test_response_body_is_valid_zip(self):
        session_obj = _make_session_with_guide()
        with patch(
            "app.services.session_service.get_session",
            new=AsyncMock(return_value=session_obj),
        ):
            client = TestClient(app, raise_server_exceptions=False)
            resp = client.get(f"/api/v1/sessions/{SESSION_ID}/export.docx")

        assert resp.status_code == 200
        assert resp.content[:2] == b"PK"

    def test_content_disposition_attachment(self):
        session_obj = _make_session_with_guide()
        with patch(
            "app.services.session_service.get_session",
            new=AsyncMock(return_value=session_obj),
        ):
            client = TestClient(app, raise_server_exceptions=False)
            resp = client.get(f"/api/v1/sessions/{SESSION_ID}/export.docx")

        cd = resp.headers.get("content-disposition", "")
        assert "attachment" in cd
        assert ".docx" in cd

    def test_returns_422_when_guide_content_is_null(self):
        session_obj = _make_session_without_guide()
        with patch(
            "app.services.session_service.get_session",
            new=AsyncMock(return_value=session_obj),
        ):
            client = TestClient(app, raise_server_exceptions=False)
            resp = client.get(f"/api/v1/sessions/{SESSION_ID}/export.docx")

        assert resp.status_code == 422

    def test_returns_404_when_session_not_found(self):
        from app.services.session_service import SessionNotFound

        with patch(
            "app.services.session_service.get_session",
            new=AsyncMock(side_effect=SessionNotFound()),
        ):
            client = TestClient(app, raise_server_exceptions=False)
            resp = client.get(f"/api/v1/sessions/{SESSION_ID}/export.docx")

        assert resp.status_code == 404
